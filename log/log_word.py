import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from report import load_data
from tools import translate
from docx import Document
from docx.shared import Pt
# from docxcompose.composer import Composer
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx
import os
from PIL import Image, ExifTags
import requests
from io import BytesIO
import env


class LogWordApp(tk.Tk):
    def __init__(self, master=None, campaign=None,agency=None, photos_path="", campaign_number="XX", output_path=""):
        self.master = master
        tk.Label(master=self.master, text="    ").grid(row=0,column=0)
        tk.Label(master=self.master,text=f"Création du rapport d'expérimentation pour {campaign}").grid(row=1,column=1, sticky="w")
        self.progressbar_element = ttk.Progressbar(master=self.master,orient="horizontal",length=350,mode="determinate")
        self.progressbar_element.grid(row=2,column=1, sticky="w")
        self.progressbar_element["value"] = 0
        self.label_text = tk.Label(master=self.master)
        self.label_text.grid(row=3,column=1,sticky="w")
        self.quit_window = False
        self.main(campaign, campaign_number, agency, photos_path, output_path)
    
    @property
    def text(self):
        """
        Read/write the line logged in the log window
        """
        return self.text_value.get()

    @text.setter
    def text(self, value):
        self.label_text.config(text = value)
        self.label_text.update()


    @property
    def progressbar(self):
        """
        Read/write the last line logged in the log window
        """
        return self.progressbar_element["value"]

    @progressbar.setter
    def progressbar(self, value):
        self.progressbar_element["value"] = value
        self.progressbar_element.update()

    def kill_window(self):
        self.master.destroy()

    def recuperation_photo(self,reference, path_photo, path_ressources):
        '''Fonction de récupération des photos à partir de path_photo, et de la référence du point de mesure
        Path_ressources est initialisé dans word_main, il précise un fichier contenant "carre_blanc.png" qui permet de garder la mise en page même en l'absence de photo
        Attention il faut que les noms de photos soit de la bonne nomenclature (ex : "stepXX_PDA1_AG-003-01-01-01_Amont_5445664_65454.pngstepXX_PDA1_AG-003-01-01-01_Amont_5445664_65454.png")
        avec en 3e position (séparée par des underscores) la référence du point de mesure, et en 4e le type de photo (Amont, Aval, Panorama, Zoom).
        Retourne un dictionnaire avec comme clé le type de photo et en valeur le path de la photo'''
        prefixe = path_photo + "/" + reference
        list_type = ["amont", "aval", "zoom", "panorama"]
        dico_nom = {}
        try:
            filenames = os.listdir(prefixe)
        except FileNotFoundError:
            for typ in list_type:
                dico_nom[typ] = path_ressources + "/carre_blanc.jpg"
        else:
            for elt in filenames:
                l_nom = elt.split("_")
                type_photo = l_nom[3].lower()
                dico_nom[type_photo] = prefixe + "/" + elt

            for typ in list_type:
                try:
                    dico_nom[typ]
                except KeyError:
                    dico_nom[typ] = path_ressources + "/carre_blanc.jpg"
        return dico_nom


    def rotation_image(self, path_photo):
        '''Fonction pour tourner les images dans le fichier word, car la rotation dans l'explorateur windows est enregistré en metadonnée, on applique cette métadonnée
        à l'image elle même dans cette fonction 
        Ne retourne rien, modifie en dur la rotation de l'image'''
        try:
            image = Image.open(path_photo)
            for orientation in ExifTags.TAGS.keys():
                if ExifTags.TAGS[orientation] == 'Orientation':
                    break
            exif = dict(image._getexif().items())
            if exif[orientation] == 3:
                image = image.rotate(180, expand=True)
            elif exif[orientation] == 6:
                image = image.rotate(270, expand=True)
            elif exif[orientation] == 8:
                image = image.rotate(90, expand=True)
            image.save(path_photo)
            image.close()

        except (AttributeError, KeyError, IndexError):
            # cases: image don't have getexif
            pass


    def main(self, campaign, num_campaign, agence, photos_path, output_path):

        self.text = 'Chargement des données...'
        path_ressources = "Ressources/"
        doc = Document(path_ressources + 'Page_de_garde.docx')
        style = doc.styles['Normal']
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)
        place_dict, year, week_start_number, week_end_number = load_data(
            campaign)

        ## Table des matières
        doc.add_page_break()
        summary = doc.add_paragraph('Table des matières')
        summary.alignment = 1
        summary.paragraph_format.space_after = Pt(5)
        summary.bold = True

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Clique droit pour mettre à jour les champs"
        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)
        p_element = paragraph._p
        
        ## Synthèse d'opération de terrain
        doc.add_page_break()
        synthesis_table_title = doc.add_heading("Synthèse d'opération de terrain")
        synthesis_table_title.alignment = 1
        synthesis_table_title.space_after = Pt(5)
        synthesis_table = doc.add_table(rows=13, cols=2, style='Style1')
        for row in range(13):
            paragraph = synthesis_table.cell(row, 1).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(10)
            paragraph.paragraph_format.space_before = Pt(10)
        
        paragraph = synthesis_table.cell(0, 0).paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.space_before = Pt(10)

        synthesis_table.cell(0,1).merge(synthesis_table.cell(0,0)).paragraphs[0].add_run(f"Campagne {num_campaign} - {year}").bold = True
        width_total = synthesis_table.cell(0,0).width
        for cell in synthesis_table.column_cells(0):
            cell.width = width_total*0.3
        for cell in synthesis_table.column_cells(1):
            cell.width = width_total

        synthesis_table.cell(0,1).paragraphs[0].alignment = 1

        synthesis_table.cell(1,0).paragraphs[0].add_run("Nombre de stations").bold = True
        synthesis_table.cell(1,1).paragraphs[0].add_run(f"{len(list(place_dict.keys()))}").alignment = 1
        
        synthesis_table.cell(2,0).paragraphs[0].add_run("Date d'exposition \n(début/fin)").bold = True
        synthesis_table.cell(2,1).paragraphs[0].add_run(f"Semaine {week_start_number} / Semaine {week_end_number}").alignment = 1
        
        synthesis_table.cell(3,0).paragraphs[0].add_run("Bioessai(s)").bold = True
        synthesis_table.cell(3,1).paragraphs[0].add_run("Bioaccumulation selon la norme Afnor-XP-T90-721").alignment = 1
        
        synthesis_table.cell(4,0).merge(synthesis_table.cell(5,0).merge(synthesis_table.cell(6,0))).paragraphs[0].add_run("Domaine d'application").bold = True

        not_conform_list = []
        conductivity_not_conform_number = 0
        temperature_not_conform_number = 0
        ph_not_conform_number = 0
        oxygen_not_conform_number = 0
        not_conform_explination_dict = {}
        for place_id in place_dict:
            not_conform_explination = " "
            if len(place_dict[place_id]["not conform"]):
                not_conform_type_list = []
                if agence:
                    not_conform_list.append((place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"]))
                else :
                    not_conform_list.append("Point " + str(place_dict[place_id]["number"]).replace(',', '-') + " : " + translate(place_dict[place_id]['name']))

                if "min_conductivity_chemistry" in place_dict[place_id]["not conform"] or "max_conductivity_chemistry" in place_dict[place_id]["not conform"] or "min_conductivity_tox" in place_dict[place_id]["not conform"] or "max_conductivity_tox" in place_dict[place_id]["not conform"]:
                    not_conform_type_list.append("conductivité")
                    conductivity_not_conform_number +=1
                    if "min_conductivity_chemistry" in place_dict[place_id]["not conform"] or "min_conductivity_tox" in place_dict[place_id]["not conform"] :
                        if "min_conductivity_chemistry" in place_dict[place_id]["not conform"] and not "min_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est inférieure aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (min. 50 µS/cm). "
                        if not "min_conductivity_chemistry" in place_dict[place_id]["not conform"] and "min_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est inférieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (min. 100 µS/cm). "
                        if "min_conductivity_chemistry" in place_dict[place_id]["not conform"] and "min_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est inférieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (min. 50 µS/cm). "
                    if "max_conductivity_chemistry" in place_dict[place_id]["not conform"] or "max_conductivity_tox" in place_dict[place_id]["not conform"] :
                        if "max_conductivity_chemistry" in place_dict[place_id]["not conform"] and not "max_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est supérieure aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (max. 2000 µS/cm). "
                        if not "max_conductivity_chemistry" in place_dict[place_id]["not conform"] and "max_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est supérieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (max. 1000 µS/cm). "
                        if "max_conductivity_chemistry" in place_dict[place_id]["not conform"] and "max_conductivity_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La conductivité est supérieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (min. 2000 µS/cm). "
                if "min_temperature_chemistry" in place_dict[place_id]["not conform"] or "max_temperature_chemistry" in place_dict[place_id]["not conform"] or "min_temperature_alimentation" in place_dict[place_id]["not conform"] or "max_temperature_alimentation" in place_dict[place_id]["not conform"] or "min_temperature_reproduction" in place_dict[place_id]["not conform"] or "max_temperature_reproduction" in place_dict[place_id]["not conform"]:
                    not_conform_type_list.append("température")
                    temperature_not_conform_number +=1
                    if "min_temperature_chemistry" in place_dict[place_id]["not conform"] or "min_temperature_alimentation" in place_dict[place_id]["not conform"]  :
                        if "min_temperature_chemistry" in place_dict[place_id]["not conform"] and not "min_temperature_alimentation" in place_dict[place_id]["not conform"] and not "min_temperature_reproduction" in place_dict[place_id]["not conform"]  :
                            not_conform_explination += f"La température est inférieure aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (mesurée {round(place_dict[place_id]['condition']['chemistry_average_temperature_min'],2)} pour min. 1 °C). "
                        if not "min_temperature_chemistry" in place_dict[place_id]["not conform"] and "min_temperature_alimentation" in place_dict[place_id]["not conform"] and not "min_temperature_reproduction" in place_dict[place_id]["not conform"]:
                            not_conform_explination += f"La température est inférieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (mesurée {round(place_dict[place_id]['condition']['alimentation_average_temperature_min'],2)} pour min. 7 °C). "
                        if not "min_temperature_chemistry" in place_dict[place_id]["not conform"] and not "min_temperature_alimentation" in place_dict[place_id]["not conform"] and "min_temperature_reproduction" in place_dict[place_id]["not conform"]:
                            not_conform_explination += f"La température est inférieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (mesurée {round(place_dict[place_id]['condition']['reproduction_average_temperature_min'],2)} pour min. 7 °C). "
                        if "min_temperature_chemistry" in place_dict[place_id]["not conform"] and "min_temperature_alimentation" in place_dict[place_id]["not conform"] and "min_temperature_reproduction" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La température est inférieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (mesurée {round(place_dict[place_id]['condition']['average_temperature'],2)} pour min. 1 °C). "
                    if "max_temperature_chemistry" in place_dict[place_id]["not conform"] or "max_temperature_alimentation" in place_dict[place_id]["not conform"]  :
                        if "max_temperature_chemistry" in place_dict[place_id]["not conform"] and not "max_temperature_alimentation" in place_dict[place_id]["not conform"] and not "max_temperature_reproduction" in place_dict[place_id]["not conform"]  :
                            not_conform_explination += f"La température est supérieure aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (mesurée {round(place_dict[place_id]['condition']['chemistry_average_temperature_max'],2)} pour max. 18 °C). "
                        if not "max_temperature_chemistry" in place_dict[place_id]["not conform"] and "max_temperature_alimentation" in place_dict[place_id]["not conform"] and not "max_temperature_reproduction" in place_dict[place_id]["not conform"]:
                            not_conform_explination += f"La température est supérieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (mesurée {round(place_dict[place_id]['condition']['alimentation_average_temperature_max'],2)} pour max. 20 °C). "
                        if not "max_temperature_chemistry" in place_dict[place_id]["not conform"] and not "max_temperature_alimentation" in place_dict[place_id]["not conform"] and "max_temperature_reproduction" in place_dict[place_id]["not conform"]:
                            not_conform_explination += f"La température est supérieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (mesurée {round(place_dict[place_id]['condition']['reproduction_average_temperature_max'],2)} pour max. 20 °C). "
                        if "max_temperature_chemistry" in place_dict[place_id]["not conform"] and "max_temperature_alimentation" in place_dict[place_id]["not conform"] and "max_temperature_reproduction" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"La température est supérieure aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (mesurée {round(place_dict[place_id]['condition']['average_temperature'],2)} pour max. 20 °C). "
                if "min_ph_chemistry" in place_dict[place_id]["not conform"] or "max_ph_chemistry" in place_dict[place_id]["not conform"] or "min_ph_tox" in place_dict[place_id]["not conform"] or "max_ph_tox" in place_dict[place_id]["not conform"]:
                    not_conform_type_list.append("ph")
                    ph_not_conform_number +=1
                    if "min_ph_chemistry" in place_dict[place_id]["not conform"] or "min_ph_tox" in place_dict[place_id]["not conform"] :
                        if "min_ph_chemistry" in place_dict[place_id]["not conform"] and not "min_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est inférieur aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (min. 6,3 µS/cm). "
                        if not "min_ph_chemistry" in place_dict[place_id]["not conform"] and "min_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est inférieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (min. 6,3 µS/cm). "
                        if "min_ph_chemistry" in place_dict[place_id]["not conform"] and "min_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est inférieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (min. 6,3 µS/cm). "
                    if "max_ph_chemistry" in place_dict[place_id]["not conform"] or "max_ph_tox" in place_dict[place_id]["not conform"] :
                        if "max_ph_chemistry" in place_dict[place_id]["not conform"] and not "max_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est supérieur aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (max. 8,9 µS/cm). "
                        if not "max_ph_chemistry" in place_dict[place_id]["not conform"] and "max_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est supérieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (max. 8,9 µS/cm). "
                        if "max_ph_chemistry" in place_dict[place_id]["not conform"] and "max_ph_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le pH est supérieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (min. 8,9 µS/cm). "
                if "oxygen_chemistry" in place_dict[place_id]["not conform"] or "oxygen_tox" in place_dict[place_id]["not conform"]:
                    not_conform_type_list.append("oxygène")
                    oxygen_not_conform_number +=1
                    if "oxygen_chemistry" in place_dict[place_id]["not conform"] or "oxygen_tox" in place_dict[place_id]["not conform"] :
                        if "oxygen_chemistry" in place_dict[place_id]["not conform"] and not "oxygen_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le taux d'oxygène est inférieur aux recommandations pour l’encagement de gammares dans le cadre du biotest de chimie (min. 5 mg/L). "
                        if not "oxygen_chemistry" in place_dict[place_id]["not conform"] and "oxygen_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le taux d'oxygène est inférieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de toxicité (min. 5 mg/L). "
                        if "oxygen_chemistry" in place_dict[place_id]["not conform"] and "oxygen_tox" in place_dict[place_id]["not conform"] :
                            not_conform_explination += f"Le taux d'oxygène est inférieur aux recommandations pour l’encagement de gammares dans le cadre des biotests de chimie et de toxicité (min. 5 mg/L). "
                precision = ' ('
                for index, not_conform_type in enumerate(not_conform_type_list):
                    if index==len(not_conform_type_list)-1 and index>0 :
                        precision+=" et " + not_conform_type + ")"
                    elif index==len(not_conform_type_list)-1 :
                        precision+=not_conform_type + ")"
                    elif index==0 :
                        precision+=not_conform_type
                    else :
                        precision+=", " + not_conform_type
                not_conform_list[-1]+=precision
            not_conform_explination_dict[place_id] = not_conform_explination
        synthesis_table.cell(4,1).paragraphs[0].add_run("Conforme\n").bold = True
        synthesis_table.cell(4,1).paragraphs[0].add_run(f"{len(list(place_dict.keys())) - len(not_conform_list)} stations sur {len(list(place_dict.keys()))} ")
        synthesis_table.cell(5,1).paragraphs[0].add_run("Non Conforme\n").bold = True
        not_conform_resume = "Aucune station"
        if len(not_conform_list):
            if len(not_conform_list) == 1:
                not_conform_resume = "1 station (voir ci-dessous)\n"
                if conductivity_not_conform_number:
                    not_conform_resume+="Conductivité" 
                if temperature_not_conform_number:
                    not_conform_resume+="Température" 
                if ph_not_conform_number:
                    not_conform_resume+="pH"
            else :
                not_conform_resume = f"{len(not_conform_list)} stations (voir ci-dessous)\n" 
                if conductivity_not_conform_number:
                    not_conform_resume+=f"Conductivité({conductivity_not_conform_number})"
                if temperature_not_conform_number:
                    not_conform_resume+=(", " if conductivity_not_conform_number else "") + f"Température({temperature_not_conform_number})"
                if ph_not_conform_number:
                    not_conform_resume+=(", " if conductivity_not_conform_number or temperature_not_conform_number else "") + f"pH({ph_not_conform_number})"

        synthesis_table.cell(5,1).paragraphs[0].add_run(not_conform_resume)
        for not_conform_station in not_conform_list:
            synthesis_table.cell(6,1).paragraphs[0].add_run(not_conform_station + "\n")
        
        vandalism_list = []
        for place_id in place_dict:
            if "vandalism" in place_dict[place_id]:
                vandalism_list.append(place_id)
        not_validated_list = []
        for place_id in place_dict:
            if "chemistry portion validation" in place_dict[place_id] and place_dict[place_id]["chemistry portion validation"]==0 and place_id not in vandalism_list and place_dict[place_id]["chemistry_survival"]!="0%":
                if agence:
                    not_validated_list.append((place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"]))
                else :
                    not_validated_list.append("Point " + str(place_dict[place_id]["number"]).replace(',', '-') + " : " + translate(place_dict[place_id]['name']))

        synthesis_table.cell(7,0).merge(synthesis_table.cell(8,0).merge(synthesis_table.cell(9,0))).paragraphs[0].add_run("Prise d’essai minimale pour la réalisation de l’intégralité des analyses chimiques*").bold = True
        synthesis_table.cell(7,1).paragraphs[0].add_run("Atteinte\n").bold = True
        synthesis_table.cell(7,1).paragraphs[0].add_run(f"{len(list(place_dict.keys())) - len(not_validated_list)} stations sur {len(list(place_dict.keys()))} ")
        synthesis_table.cell(8,1).paragraphs[0].add_run("Non atteinte\n").bold = True
        not_validate_resume = "Aucune station"
        if len(not_validated_list):
            if len(not_validated_list) == 1:
                not_validate_resume = "1 station (voir ci-dessous)"
            else :
                not_validate_resume = f"{len(not_validated_list)} stations (voir ci-dessous)" 
                
        synthesis_table.cell(8,1).paragraphs[0].add_run(not_validate_resume)
        for not_validated_station in not_validated_list:
            synthesis_table.cell(9,1).paragraphs[0].add_run(not_validated_station + "\n")
        synthesis_table.cell(10,0).paragraphs[0].add_run("Vandalisme/Perte").bold = True
        for index, place_id in enumerate(vandalism_list):
            if agence:
                place_description = (place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"])
            else :
                place_description = "Point " + str(place_dict[place_id]["number"]).replace(',', '-') + " : " + translate(place_dict[place_id]['name'])
            if index>0 :
                synthesis_table.cell(10,1).paragraphs[0].add_run("\n" + place_description).bold = True
            else :    
                synthesis_table.cell(10,1).paragraphs[0].add_run(place_description)

        synthesis_table.cell(11,0).paragraphs[0].add_run("Survie nulle \n(Bioaccumulation)").bold = True
        null_survival = ""
        for place_id in place_dict:
            if "chemistry_survival" in place_dict[place_id] and place_dict[place_id]["chemistry_survival"]=="0%":
                if null_survival=="":
                    null_survival= (place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"])
                else :
                    null_survival="\n" + (place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"])
        
        synthesis_table.cell(11,1).paragraphs[0].add_run(null_survival if null_survival!="" else "Aucune station")
        synthesis_table.cell(12,0).paragraphs[0].add_run("Autres remarques").bold = True
        doc.add_paragraph().add_run("*La prise d’essai est confirmée par le laboratoire d’analyses chimiques après lyophilisation de l’échantillon. En cas de non atteinte de la prise d’essai minimale, le laboratoire peut demander au client de faire certains choix de techniques d’analyses.").italic=True

        length = len(list(place_dict.keys()))

        self.progressbar_element["maximum"] = length+1

        count=1
        self.progressbar +=1
        ## Stations de mesure
        for place_id in place_dict:
            self.text = f"Création de la page {count}/{length}"
            doc.add_page_break()
            if agence and "reference" in place_dict[place_id]:
                title = doc.add_heading((place_dict[place_id]["agency"] +" : " if "agency" in place_dict[place_id] else "")+ translate(place_dict[place_id]["name"]) + "   " + place_dict[place_id]["reference"] )
                title.alignment = 1
                title.bold = True
                title.space_after = Pt(5)
            else:
                title = doc.add_heading("Point " + str(place_dict[place_id]["number"]).replace(',', '-') + " : " + translate(place_dict[place_id]['name']))
                title.alignment = 1
                title.bold = True
                title.space_after = Pt(5)
                
            table_geo_1 = doc.add_table(rows=2, cols=6)
            
            table_geo_1.cell(0, 1).merge(table_geo_1.cell(0, 2))
            table_geo_1.cell(0, 4).merge(table_geo_1.cell(0, 5))
            table_geo_1.cell(1, 0).merge(table_geo_1.cell(1, 1))
            table_geo_1.cell(1, 2).merge(table_geo_1.cell(1, 5))

            if agence:
                table_geo_2 = doc.add_table(rows=5, cols=4)
                for j in range(0, 5):
                    table_geo_2.cell(j, 0).merge(table_geo_2.cell(j, 1))
                for j in range(0, 1):
                    table_geo_2.cell(j, 2).merge(table_geo_2.cell(j, 3))
            else:
                table_geo_2 = doc.add_table(rows=1, cols=4)
                table_geo_2.cell(0, 0).merge(table_geo_2.cell(0, 1))


            width_table = table_geo_1.cell(0, 1).width

            table_geo_1.cell(0, 0).paragraphs[0].add_run(
                'Commune :').bold = True
            table_geo_1.cell(0, 1).paragraphs[0].add_run(
                translate(place_dict[place_id]['city']) if 'city' in place_dict[place_id] else ''  + "    " + place_dict[place_id]['zipcode'] if 'zipcode' in place_dict[place_id] else '')
            table_geo_1.cell(0, 4).paragraphs[0].add_run(
                translate(place_dict[place_id]['stream']) if 'stream' in place_dict[place_id] else '')

            table_geo_1.cell(0, 3).paragraphs[0].add_run(
                "Cours d'eau : ").bold = True

            table_geo_1.cell(1, 0).paragraphs[0].add_run(
                "Biotests :").bold = True
            biotest_list = []
            for measurepoint_id in place_dict[place_id]["measurepoint"]:
                for pack_id in place_dict[place_id]["measurepoint"][measurepoint_id]["pack"]:
                    if place_dict[place_id]["measurepoint"][measurepoint_id]["pack"][pack_id] not in biotest_list :
                        biotest_list.append(place_dict[place_id]["measurepoint"][measurepoint_id]["pack"][pack_id])
            for index, biotest in enumerate(biotest_list):
                if biotest == "alimentation":
                    biotest_list[index] = "Alimentation"
                if biotest == "chemistry":
                    biotest_list[index] = "Chimie"
                if biotest == "reproduction":
                    biotest_list[index] = "Reproduction"
                if biotest == "neurology":
                    biotest_list[index] = "Neurotoxicité"
            french_names_biotest = None
            for biotest in biotest_list :
                if french_names_biotest:
                    french_names_biotest +=", " + biotest
                else :
                    french_names_biotest = biotest
            table_geo_1.cell(1, 2).paragraphs[0].add_run(
                french_names_biotest)

            if agence:
                table_geo_2.cell(0, 0).paragraphs[0].add_run(
                    "Réseau de surveillance :").bold = True
                table_geo_2.cell(0, 2).paragraphs[0].add_run(
                    translate(place_dict[place_id]['network']) if "network" in place_dict[place_id] else "ND" )

                table_geo_2.cell(1, 0).paragraphs[0].add_run(
                    "Type d'hydroécorégion :").bold = True
                table_geo_2.cell(1, 2).paragraphs[0].add_run(
                    translate(place_dict[place_id]['hydroecoregion']) if "hydroecoregion" in place_dict[place_id] else "ND")

                table_geo_2.cell(2, 0).paragraphs[0].add_run(
                    "Coordonnées Agence Lambert 93 :").bold = True
                table_geo_2.cell(2, 2).paragraphs[0].add_run('Y ' + place_dict[place_id]['lambertY'] if "lambertY" in place_dict[place_id] else "ND")
                table_geo_2.cell(2, 3).paragraphs[0].add_run('X ' + place_dict[place_id]['lambertX'] if "lambertX" in place_dict[place_id] else "ND")

                table_geo_2.cell(3, 0).paragraphs[0].add_run("Coordonnées BIOMÆ en degrés décimaux : ").bold = True
                longitude = place_dict[place_id]['longitudeSpotted']
                latitude = place_dict[place_id]['latitudeSpotted']
                table_geo_2.cell(3, 2).paragraphs[0].add_run(
                    str(longitude))
                table_geo_2.cell(3, 3).paragraphs[0].add_run(
                    str(latitude))

                table_geo_2.cell(4, 0).paragraphs[0].add_run(
                    "Coordonnées BIOMÆ Lambert 93 : ").bold = True
                table_geo_2.cell(4, 2).paragraphs[0].add_run('Y ' + place_dict[place_id]['lambertYSpotted'].replace('.', ','))
                table_geo_2.cell(4, 3).paragraphs[0].add_run('X ' + place_dict[place_id]['lambertXSpotted'].replace('.', ','))

            else:
                table_geo_2.cell(0, 0).paragraphs[0].add_run(
                    "Coordonnées BIOMÆ en degrés décimaux : ").bold = True
                longitude = place_dict[place_id]['longitudeSpotted']
                latitude = place_dict[place_id]['latitudeSpotted']
                table_geo_2.cell(0, 2).paragraphs[0].add_run(
                    str(place_dict[place_id]['longitudeSpotted']))
                table_geo_2.cell(0, 3).paragraphs[0].add_run(
                    str(place_dict[place_id]['latitudeSpotted']))

            table_carte = doc.add_table(rows=4, cols=1)
            lon = str(place_dict[place_id]['longitudeSpotted'])
            lat = str(place_dict[place_id]['latitudeSpotted'])
            if (lon != "None") & (lat != "None"):
                access_token = env.ACCESS_TOKEN_MAPBOX
                # layer = '{"id":"water","source":{"url":"mapbox://mapbox.mapbox-streets-v8","type":"vector"},"source-layer":"water","type":"fill","paint":{"fill-color":"%2300ffff"}}'
                url_street = f"https://api.mapbox.com/styles/v1/mapbox/streets-v11/static/pin-s+FF0000({lon},{lat})/{lon},{lat},9.21/450x300@2x?access_token={access_token}"
                response = requests.get(url_street)
                carte_street = BytesIO(response.content)
                table_carte.cell(0, 0).paragraphs[0].add_run().add_picture(
                    carte_street, width=4500000)
            else:
                table_carte.cell(0, 0).paragraphs[0].add_run().add_picture(
                    path_ressources + "/carre_blanc.jpg", width=4500000)
            table_carte.cell(0, 0).paragraphs[0].alignment = 1
            table_carte.cell(1, 0).text = "Localisation du point de mesure"
            table_carte.cell(1, 0).paragraphs[0].alignment = 1

            if (lon != "None") & (lat != "None"):
                url_satellite = f"https://api.mapbox.com/styles/v1/mapbox/satellite-streets-v11/static/pin-s+FF0000({lon},{lat})/{lon},{lat},13.5/450x300@2x?access_token={access_token}"
                response = requests.get(url_satellite)
                carte_satellite = BytesIO(response.content)
                table_carte.cell(2, 0).paragraphs[0].add_run().add_picture(
                    carte_satellite, width=4500000)
            else:
                table_carte.cell(2, 0).paragraphs[0].add_run().add_picture(
                    path_ressources + "/carre_blanc.jpg", width=4500000)
            table_carte.cell(2, 0).paragraphs[0].alignment = 1
            table_carte.cell(3, 0).text = "Vue satellitaire"
            table_carte.cell(3, 0).paragraphs[0].alignment = 1

            doc.add_page_break()

            table_image = doc.add_table(rows=8, cols=2)
            table_image.cell(0, 0).merge(table_image.cell(0, 1))
            table_image.cell(1, 0).merge(table_image.cell(1, 1))

            if agence and place_dict[place_id]:
                table_image.cell(0, 0).paragraphs[0].add_run((place_dict[place_id]['agency']+ " : " if 'agency' in place_dict[place_id] else "") + translate(place_dict[place_id]['name']) + "  " + place_dict[place_id]["reference"]).bold = True
            else:
                table_image.cell(0, 0).paragraphs[0].add_run("Point " + str(place_dict[place_id]["number"]).replace(',','-') + " : " +
                                                                translate(place_dict[place_id]['name'])).bold = True
            table_image.cell(0, 0).paragraphs[0].alignment = 1
            table_image.cell(
                0, 0).paragraphs[0].paragraph_format.line_spacing = font.size

            table_image.cell(1, 0).paragraphs[0].add_run("Photos de la station de mesure de la qualité des eaux pour la campagne " + num_campaign + "-" + str(year)).bold = True  # Mettre que l'année, passage en argument ou autre méthode de récupération ?
            table_image.cell(1, 0).paragraphs[0].alignment = 1
            table_image.cell(1, 0).paragraphs[0].paragraph_format.line_spacing = font.size

            table_image.cell(3, 0).text = "Aval de zone d’encagement"
            table_image.cell(3, 0).paragraphs[0].alignment = 1
            table_image.cell(3, 1).text = "Amont de zone d’encagement"
            table_image.cell(3, 1).paragraphs[0].alignment = 1
            table_image.cell(5, 0).text = "Gros plan de l’encagement"
            table_image.cell(5, 0).paragraphs[0].alignment = 1
            table_image.cell(5, 1).text = "Panorama encagement"
            table_image.cell(5, 1).paragraphs[0].alignment = 1

            for row in range(3, 8):
                if row != 4:
                    table_image.cell(row, 0).paragraphs[0].paragraph_format.line_spacing = font.size
                    table_image.cell(row, 1).paragraphs[0].paragraph_format.line_spacing = font.size

            nom_photo = self.recuperation_photo(place_dict[place_id]["reference_photo"], photos_path, path_ressources)
            self.rotation_image(nom_photo['amont'])
            self.rotation_image(nom_photo['aval'])
            self.rotation_image(nom_photo['zoom'])
            self.rotation_image(nom_photo['panorama'])

            table_image.cell(2, 0).paragraphs[0].add_run().add_picture(
                nom_photo['aval'], width=3046870)  # width=3046870, height=2111370
            table_image.cell(2, 1).paragraphs[0].add_run().add_picture(
                nom_photo['amont'], width=3046870)
            table_image.cell(4, 0).paragraphs[0].add_run().add_picture(
                nom_photo['zoom'], width=3046870)
            table_image.cell(4, 1).paragraphs[0].add_run().add_picture(
                nom_photo['panorama'], width=3046870)
            for elt in [(2, 0), (2, 1), (4, 0), (4, 1)]:
                table_image.cell(elt[0], elt[1]).paragraphs[0].paragraph_format.space_after = Pt(0)
                table_image.cell(elt[0], elt[1]).paragraphs[0].paragraph_format.space_before = Pt(0)

            table_image.cell(6, 0).paragraphs[0].add_run("Type de système d’exposition : ").bold = True

            type_barrel_J0 = place_dict[place_id]["barrel_type"] if "barrel_type" in place_dict[place_id] else ""
            if (type_barrel_J0 == 'barrel'):
                type_barrel_J0 = 'Fut'
            elif (type_barrel_J0 == 'box'):
                type_barrel_J0 = 'Caisse'
            table_image.cell(6, 1).paragraphs[0].add_run(type_barrel_J0)
            table_image.cell(7, 0).merge(table_image.cell(7, 1))
            table_image.cell(7, 0).paragraphs[0].add_run("Paramètres physico-chimiques pour la campagne : " + num_campaign + "-" + str(year)).bold = True

            table_temperature = doc.add_table(rows=2, cols=4, style="Table Grid")
            table_temperature.cell(0, 0).merge(table_temperature.cell(1, 0))
            table_temperature.cell(0, 0).paragraphs[0].add_run("Température eau (°C) Sonde en continu").italic = True
            table_temperature.cell(0, 0).paragraphs[0].alignment = 1

            table_temperature.cell(0, 1).paragraphs[0].add_run("Minimum")  # .bold = True
            table_temperature.cell(0, 1).paragraphs[0].alignment = 1
            table_temperature.cell(0, 2).paragraphs[0].add_run("Moyenne")  # .bold = True
            table_temperature.cell(0, 2).paragraphs[0].alignment = 1
            table_temperature.cell(0, 3).paragraphs[0].add_run("Maximum")  # .bold = True
            table_temperature.cell(0, 3).paragraphs[0].alignment = 1
            if place_dict[place_id]["condition"]['temperature_min'] is not None:
                table_temperature.cell(1, 1).paragraphs[0].add_run(str(round(place_dict[place_id]["condition"]['temperature_min'], 1)).replace('.', ','))
            table_temperature.cell(1, 1).paragraphs[0].alignment = 1
            if place_dict[place_id]["condition"]['average_temperature'] is not None:
                table_temperature.cell(1, 2).paragraphs[0].add_run(str(round(place_dict[place_id]["condition"]['average_temperature'], 1)).replace('.', ','))
            table_temperature.cell(1, 2).paragraphs[0].alignment = 1
            if place_dict[place_id]["condition"]['temperature_max'] is not None:
                table_temperature.cell(1, 3).paragraphs[0].add_run(str(round(place_dict[place_id]["condition"]['temperature_max'], 1)).replace('.', ','))
            table_temperature.cell(1, 3).paragraphs[0].alignment = 1
            for row in range(2):
                for col in range(4):
                    paragraph = table_temperature.cell(row, col).paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.space_before = Pt(4)

            interligne = doc.add_paragraph()
            interligne.paragraph_format.space_after = Pt(0)
            interligne.paragraph_format.space_before = Pt(0)

            chemistry = 1 if "Chimie" in biotest_list else 0
            usefull_days = []
            for J in place_dict[place_id]["condition"]:
                if place_dict[place_id]["condition"][J]:
                    if J == 'J0':
                        usefull_days.append(['J0','J+0', place_dict[place_id]["condition"][J]["date"] if "date" in place_dict[place_id]["condition"][J] else None,place_dict[place_id]["condition"][J]["comment"] if "comment" in place_dict[place_id]["condition"][J] else None])
                    if J == 'J7':
                        usefull_days.append(['J7','J+7', place_dict[place_id]["condition"][J]["date"] if "date" in place_dict[place_id]["condition"][J] else None,place_dict[place_id]["condition"][J]["comment"] if "comment" in place_dict[place_id]["condition"][J] else None])
                    if J == 'J14':
                        usefull_days.append(['J14','J+14', place_dict[place_id]["condition"][J]["date"] if "date" in place_dict[place_id]["condition"][J] else None,place_dict[place_id]["condition"][J]["comment"] if "comment" in place_dict[place_id]["condition"][J] else None])
                    if J == 'J21':
                        usefull_days.append(['J21','J+21', place_dict[place_id]["condition"][J]["date"] if "date" in place_dict[place_id]["condition"][J] else None,place_dict[place_id]["condition"][J]["comment"] if "comment" in place_dict[place_id]["condition"][J] else None])
                    if J == 'JN':
                        usefull_days.append(['JN','J+N', place_dict[place_id]["condition"][J]["date"] if "date" in place_dict[place_id]["condition"][J] else None,place_dict[place_id]["condition"][J]["comment"] if "comment" in place_dict[place_id]["condition"][J] else None])

            table_exposure_condition = doc.add_table(rows=7+chemistry, cols=1+len(usefull_days), style="Table Grid")
            liste_entete = ["Intervention", "Date - Heure","Température (°C)", "Conductivité (µS/cm)", "pH", "Oxygène dissous (mg/L)", "Survie Chimie (%)"]
            liste_entete_BDD = ["date", "temperature","conductivity", "ph", "oxygen"]
            for num_entete in range(6+chemistry):
                paragraph = table_exposure_condition.cell(num_entete, 0).paragraphs[0]
                paragraph.add_run(liste_entete[num_entete]).italic = True
                paragraph.alignment = 1
                table_exposure_condition.cell(num_entete, 0).width = width_table*0.8
            for num_jour in range(len(usefull_days)):
                paragraph = table_exposure_condition.cell(0, num_jour+1).paragraphs[0]
                paragraph.add_run(usefull_days[num_jour][1]).bold = True
                paragraph.alignment = 1
            for num_entete in range(5):
                for num_jour in range(len(usefull_days)):
                    paragraph = table_exposure_condition.cell(num_entete+1, num_jour+1).paragraphs[0]
                    value = place_dict[place_id]["condition"][usefull_days[num_jour][0]][liste_entete_BDD[num_entete]] if liste_entete_BDD[num_entete] in place_dict[place_id]["condition"][usefull_days[num_jour][0]] else "ND"
                    if (liste_entete_BDD[num_entete] == "conductivity") & (value is not None):
                        paragraph.add_run(str(int(value) if value!='ND' else "ND"))
                    else:
                        paragraph.add_run(str(
                            value).replace('.', ','))
                    paragraph.alignment = 1

            if chemistry:
                table_exposure_condition.cell(6, 1).merge(
                    table_exposure_condition.cell(6, len(usefull_days)))
                paragraph_survie = table_exposure_condition.cell(
                    6, 1).paragraphs[0]
                paragraph_survie.add_run((place_dict[place_id]['chemistry_survival'] if 'chemistry_survival' in place_dict[place_id] else 'ND').replace('.', ','))
                paragraph_survie.alignment = 1

                table_exposure_condition.cell(7, 0).merge(table_exposure_condition.cell(7, len(usefull_days)))
                paragraph_comment = table_exposure_condition.cell(7, 0).paragraphs[0]
            else:
                table_exposure_condition.cell(6, 0).merge(table_exposure_condition.cell(6, len(usefull_days)))
                paragraph_comment = table_exposure_condition.cell(6, 0).paragraphs[0]

            comment = ""
            for jour in usefull_days:
                if jour[3] and jour[3]!="":
                    comment +=("\n" if comment!="" else "") + f"{jour[0]} : "+ translate(jour[3].replace('\n',' '))
            comment += not_conform_explination_dict[place_id]
            if comment==" ":
                comment = "Aucune remarque particulière"
            paragraph_comment.add_run('Commentaires : \n').bold = True
            paragraph_comment.add_run(comment)
            for num_entete in range(7+chemistry):
                for num_jour in range(1+len(usefull_days)):
                    paragraph = table_exposure_condition.cell(num_entete, num_jour).paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.space_before = Pt(4)
            self.progressbar += 1
            count +=1

        doc.add_page_break()
        doc.add_paragraph().add_run().add_picture(f"{path_ressources}/Page_fin.png", width=width_total)
        doc.save(output_path)
        

        self.text = "Terminé"

        tk.Button(master=self.master, text="OK", command=self.kill_window).grid(row=4, column=1, ipadx=30)
    
