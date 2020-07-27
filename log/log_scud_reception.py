import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from report import load_data
from tools import QueryScript, translate
from docx import Document
from docx.shared import Pt, Cm
import docx
import env


class LogWordReceptionApp(tk.Tk):
    def __init__(self, master=None, campaign=None, output_path=""):
        self.master = master
        tk.Label(master=self.master, text="    ").grid(row=0,column=0)
        tk.Label(master=self.master,text=f"Création de la fiche réception de Gammares pour {campaign}").grid(row=1,column=1, sticky="w")
        self.progressbar_element = ttk.Progressbar(master=self.master,orient="horizontal",length=350,mode="determinate")
        self.progressbar_element.grid(row=2,column=1, sticky="w")
        self.progressbar_element["value"] = 0
        self.progressbar_element["maximum"]=2
        self.label_text = tk.Label(master=self.master)
        self.label_text.grid(row=3,column=1,sticky="w")
        self.quit_window = False
        self.main(campaign, output_path)
    
    @property
    def text(self):
        """
        Read/write the line logged in the log window
        """
        return self.label_text.get()

    @text.setter
    def text(self, value):
        self.label_text.config(text = value)
        self.label_text.update()


    @property
    def progressbar(self):
        """
        Read/write the last line logged in the log window
        """
        return self.progressbar_element["value"]

    @progressbar.setter
    def progressbar(self, value):
        self.progressbar_element["value"] = value
        self.progressbar_element.update()

    def kill_window(self):
        self.master.destroy()

    def create_table(self, doc, point_list):
        table = doc.add_table(rows=len(point_list)*2+1, cols=9, style="Style2")
        table.cell(0,0).paragraphs[0].add_run("Libellé de l'échantillon").alignment = 1
        table.cell(0,1).paragraphs[0].add_run("Flacon").alignment = 1
        table.cell(0,2).paragraphs[0].add_run("Qté").alignment = 1
        table.cell(0,3).paragraphs[0].add_run("Etiquette LDA").alignment = 1
        table.cell(0,4).paragraphs[0].add_run("Validation Labo").alignment = 1
        table.cell(0,5).paragraphs[0].add_run("Commentaire").alignment = 1
        table.cell(0,6).paragraphs[0].add_run("Lyophi\nlisation").alignment = 1
        table.cell(0,7).paragraphs[0].add_run("N° planification LDA").alignment = 1
        table.cell(0,8).paragraphs[0].add_run("N° LDA").alignment = 1
        for point_index, point in enumerate(point_list) :
            self.add_row(table, point, point_index)
        for row in table.rows :
            row.height = Cm(0.7)

    def add_row(self, table, point, point_row) :
        [reference, pack_comment, pack_sampling_comment, J21_comment, metal_weight, organic_weight, date] = point
        comment = (pack_comment if pack_comment else '') + '\t ' + (pack_sampling_comment if pack_sampling_comment else '') + '\t ' + (J21_comment if J21_comment else '')
            
        filtered_comment = ''
        if comment!='' :
            dissociated_comment = [translate(information).replace("\n",", ") for information in comment.split("\t")]
            need_to_be_removed = []
            for information in dissociated_comment:
                if information!="":
                    if information[-1]=="%" or "Scan" in information or "#" in information or 'RAS' in information or 'ras' in information:
                        need_to_be_removed.append(information)
                    else :
                        try :
                            int(information)
                            need_to_be_removed.append(information)
                        except ValueError :
                            pass
            for element in need_to_be_removed:
                dissociated_comment.remove(element)
            if len(dissociated_comment):
                filtered_comment = ', '.join(dissociated_comment)
        if metal_weight and metal_weight<500:
            filtered_comment += "Métaux < 500mg"
        if organic_weight and organic_weight<2500:
            filtered_comment += "Orga < 2500mg"
        
        table.cell(point_row*2+1, 0).width = self.table_width 
        table.cell(point_row*2+1, 0).merge(table.cell(point_row*2+2, 0)).paragraphs[0].add_run(reference).alignment = 1
        table.cell(point_row*2+1, 1).width = self.table_width * 0.02
        table.cell(point_row*2+1, 1).paragraphs[0].add_run("Mét")
        table.cell(point_row*2+2, 1).paragraphs[0].add_run("Orga")
        table.cell(point_row*2+1, 2).width = self.table_width * 0.02
        table.cell(point_row*2+1, 2).paragraphs[0].add_run("1" if metal_weight_valid else "0")
        table.cell(point_row*2+2, 2).paragraphs[0].add_run("1" if organic_weight_valid else "0")
        table.cell(point_row*2+1, 3).merge(table.cell(point_row*2+2, 3))
        table.cell(point_row*2+1, 3).width = self.table_width *0.8
        table.cell(point_row*2+1, 4).width = self.table_width * 0.02

        table.cell(point_row*2+1, 5).width = self.table_width * 0.6
        table.cell(point_row*2+1, 5).merge(table.cell(point_row*2+2, 5)).paragraphs[0].add_run(filtered_comment)
        table.cell(point_row*2+1, 6).width = self.table_width * 0.02
        table.cell(point_row*2+1, 6).merge(table.cell(point_row*2+2, 6))
   
    def main(self, campaign, output_path):
        
        self.text = 'Chargement des données...'
        data = QueryScript(f"SELECT Measurepoint.reference, Pack.comment, Pack.sampling_comment, MeasureExposureCondition.comment, Pack.sampling_weight-Pack.metal_tare_bottle_weight, Pack.organic_total_weight- Pack.organic_tare_bottle_weight, key_dates.date FROM {env.DATABASE_RAW}.Measurepoint JOIN {env.DATABASE_RAW}.Pack ON Pack.measurepoint_id=Measurepoint.id JOIN {env.DATABASE_TREATED}.key_dates ON key_dates.measurepoint_id=Measurepoint.id JOIN {env.DATABASE_RAW}.MeasureExposureCondition ON MeasureExposureCondition.measurepoint_id=Measurepoint.id AND key_dates.date=MeasureExposureCondition.recordedAt WHERE nature = 'chemistry' AND Measurepoint.reference like '{campaign}%' AND version={env.CHOSEN_VERSION()} AND date_id=7;").execute()
        year = data[0][-1].year
        self.progressbar += 1 
        self.text = 'Création du fichier...'
        path_ressources = "Ressources/"
        doc = Document(path_ressources + 'Fiche réception BIOMAE Gammare Vierge.docx')
        style = doc.styles['Normal']
        font = style.font
        font.name = "Arial"
        font.size = Pt(10)

        
        heading_table = doc.add_table(rows=2, cols=3, style="Style2")
        heading_table.cell(0,0).paragraphs[0].add_run("Demandeur d'analyse").bold = True
        heading_table.cell(0,0).paragraphs[0].alignment = 1
        heading_table.cell(0,1).paragraphs[0].add_run("Payeur de l'analyse (SAUF LYOPHILISATION)").bold = True
        heading_table.cell(0,1).paragraphs[0].alignment = 1

        heading_table.cell(1,0).paragraphs[0].add_run("BIOMÆ \n").bold = True
        heading_table.cell(1,0).paragraphs[0].add_run("Adresse : 320 rue de la Outarde \n 01500 Château Gaillard \nTéléphone : 04 74 61 17 42 \nEmail :")	
        
        heading_table.cell(1,1).paragraphs[0].add_run(f"Agence de l'eau {campaign.split('-')[0] +' '+ str(year)}").bold = True
        heading_table.cell(1,1).paragraphs[0].add_run("\nAdresse : \n\nTéléphone : \nEmail : \n")
        heading_table.cell(1,2).paragraphs[0].add_run(f"Bon de commande :  Liste {campaign.split('-')[0]}\nN° de devis : \nN° de sous devis : ")
        doc.add_paragraph().add_run()
        
        self.table_width = heading_table.cell(0,0).width
        self.create_table(doc, data[:3])
        
        doc.add_paragraph().add_run().add_picture(f"{path_ressources}/Precision fiche réception.PNG", width=3*self.table_width)
        doc.add_page_break()
        self.create_table(doc, data[3:])

                

        count=1
        self.progressbar +=1
        doc.save(output_path)
        

        self.text = "Terminé"

        tk.Button(master=self.master, text="OK", command=self.kill_window).grid(row=4, column=1, ipadx=30)
    
