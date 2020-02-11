from docx import Document
from report import recuperation_donnee


def create_doc(campaign):
    doc = Document('Fichiers_remplissage/Page_de_garde.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = "Arial"
    dico_exposure_condition, dico_avg_tempe, dico_geo_mp, dico_geo_agency = recuperation_donnee(
        campaign)
    nb_measurepoint = len(dico_avg_tempe)
    for i in range(1, nb_measurepoint+1):
        doc.add_page_break()
        table_geo = doc.add_table(rows=8, cols=4)

        for j in range(2, 8):
            table_geo.cell(j, 0).merge(table_geo.cell(j, 1))
        for j in range(2, 5):
            table_geo.cell(j, 2).merge(table_geo.cell(j, 3))

        header = table_geo.rows[0].cells
        header[0].merge(header[-1])
        case_header = table_geo.cell(0, 0).paragraphs[0].add_run(dico_geo_agency[i]['code'] +
                                                                 " : " + dico_geo_agency[i]['name'])
        case_header.bold = True
        case_header = table_geo.cell(0, 0).paragraphs[0].alignment = 1

        table_geo.cell(1, 0).paragraphs[0].add_run('Commune :').bold = True
        table_geo.cell(1, 1).paragraphs[0].add_run(
            dico_geo_agency[i]['city'] + "    " + dico_geo_agency[i]['zipcode'])
        table_geo.cell(1, 2).paragraphs[0].add_run(
            "Cours d'eau : ").bold = True
        table_geo.cell(1, 3).paragraphs[0].add_run(
            dico_geo_agency[i]['stream'])

        table_geo.cell(2, 0).paragraphs[0].add_run("Biotests")
        table_geo.cell(2, 0).paragraphs[0].alignment = 1
        table_geo.cell(2, 2).paragraphs[0].add_run(
            "Alimentation, Neurotoxicité, Reproduction ? Trouver data")
        table_geo.cell(2, 2).paragraphs[0].alignment = 1

        table_geo.cell(3, 0).paragraphs[0].add_run("Réseau de surveillance :")
        table_geo.cell(3, 2).paragraphs[0].add_run(
            dico_geo_agency[i]['network'])

        table_geo.cell(4, 0).paragraphs[0].add_run("Type d'hydroécorégion :")
        table_geo.cell(4, 2).paragraphs[0].add_run(
            dico_geo_agency[i]['hydroecoregion'])

        table_geo.cell(5, 0).paragraphs[0].add_run(
            "Coordonnées Agence Lambert 93 :")
        table_geo.cell(5, 2).paragraphs[0].add_run('Y ' +
                                                   dico_geo_agency[i]['lambertY'])
        table_geo.cell(5, 3).paragraphs[0].add_run('X ' +
                                                   dico_geo_agency[i]['lambertX'])

        table_geo.cell(6, 0).paragraphs[0].add_run(
            "Coordonnées BIOMÆ en degrés décimaux : ")
        table_geo.cell(6, 2).paragraphs[0].add_run(
            str(dico_geo_mp[i]['longitudeSpotted']))
        table_geo.cell(6, 3).paragraphs[0].add_run(
            str(dico_geo_mp[i]['latitudeSpotted']))

        table_geo.cell(7, 0).paragraphs[0].add_run(
            "Coordonnées BIOMÆ Lambert 93 : ")
        table_geo.cell(7, 2).paragraphs[0].add_run('Y ' +
                                                   dico_geo_mp[i]['lambertYSpotted'])
        table_geo.cell(7, 3).paragraphs[0].add_run('X ' +
                                                   dico_geo_mp[i]['lambertXSpotted'])

        doc.add_page_break()

        table_image = doc.add_table(rows=7, cols=2)
        table_image.cell(0, 0).merge(table_image.cell(0, 1))
        table_image.cell(1, 0).merge(table_image.cell(1, 1))

        table_image.cell(0, 0).paragraphs[0].add_run(dico_geo_agency[i]['code'] +
                                                     " : " + dico_geo_agency[i]['name']).bold = True
        table_image.cell(0, 0).paragraphs[0].alignment = 1

        table_image.cell(1, 0).paragraphs[0].add_run(
            "Photos de la station de mesure de la qualité des eaux pour la campagne " + campaign[-2:] + "-" + dico_exposure_condition[i]["J+0"]["date"][5:9]).bold = True  # Mettre que l'année, passage en argument ou autre méthode de récupération ?
        table_image.cell(1, 0).paragraphs[0].alignment = 1

        table_image.cell(3, 0).text = "Aval de zone d’encagement"
        table_image.cell(3, 0).paragraphs[0].alignment = 1
        table_image.cell(3, 1).text = "Amont de zone d’encagement"
        table_image.cell(3, 1).paragraphs[0].alignment = 1
        table_image.cell(5, 0).text = "Gros plan de l’encagement"
        table_image.cell(5, 0).paragraphs[0].alignment = 1
        table_image.cell(5, 1).text = "Panorama encagement"
        table_image.cell(5, 1).paragraphs[0].alignment = 1

        table_image.cell(2, 0).paragraphs[0].add_run().add_picture(
            'Fichiers_remplissage/step50_PDA1_AG-003-01-01-01_Amont_20190219_100021.jpg', width=3046870, height=2111370)
        table_image.cell(2, 1).paragraphs[0].add_run().add_picture(
            'Fichiers_remplissage/step50_PDA1_AG-003-01-01-01_Aval_20190219_095956.jpg', width=3046870, height=2111370)
        table_image.cell(4, 0).paragraphs[0].add_run().add_picture(
            'Fichiers_remplissage/step50_PDA1_AG-003-01-01-01_Zoom_20190219_101351.jpg', width=3046870, height=2111370)
        table_image.cell(4, 1).paragraphs[0].add_run().add_picture(
            'Fichiers_remplissage/step50_PDA1_AG-003-01-01-01_Panorama_20190219_101429.jpg', width=3046870, height=2111370)

        table_image.cell(6, 0).paragraphs[0].add_run(
            "Type de système d’exposition : ").bold = True
        # Faire une fonction pour dire si on met fut ou caisse en demandant des précisions
        type_barrel_J0 = dico_exposure_condition[i]['J+0']['type']
        type_barrel_J14 = dico_exposure_condition[i]['J+14']['type']
        type_barrel_J0 = type_barrel_J14 if type_barrel_J14 else type_barrel_J0
        type_barrel_J14 = type_barrel_J0 if type_barrel_J0 else type_barrel_J14
        if (type_barrel_J0 == 'barrel') & (type_barrel_J14 == 'barrel'):
            type_barrel_J0 = 'Fut'
        elif (type_barrel_J0 == 'box') & (type_barrel_J14 == 'box'):
            type_barrel_J0 = "Caisse"
        table_image.cell(6, 1).paragraphs[0].add_run(type_barrel_J0)

        table_temperature = doc.add_table(rows=3, cols=4, style="Table Grid")
        table_temperature.cell(0, 0).merge(table_temperature.cell(0, 3))
        table_temperature.cell(1, 0).merge(table_temperature.cell(2, 0))
        table_temperature.cell(0, 0).paragraphs[0].add_run(
            "Paramètres physico-chimiques pour la campagne : " + campaign[-2:] + "-" + dico_exposure_condition[i]["J+0"]["date"][5:9]).bold = True
        table_temperature.cell(1, 0).paragraphs[0].add_run(
            "Température eau (°C) Sonde en continu").italic = True
        table_temperature.cell(1, 0).paragraphs[0].alignment = 1

        table_temperature.cell(1, 1).paragraphs[0].add_run(
            "Minimum")  # .bold = True
        table_temperature.cell(1, 1).paragraphs[0].alignment = 1
        table_temperature.cell(1, 2).paragraphs[0].add_run(
            "Moyenne")  # .bold = True
        table_temperature.cell(1, 2).paragraphs[0].alignment = 1
        table_temperature.cell(1, 3).paragraphs[0].add_run(
            "Maximum")  # .bold = True
        table_temperature.cell(1, 3).paragraphs[0].alignment = 1
        table_temperature.cell(2, 1).paragraphs[0].add_run(str(round(
            dico_avg_tempe[i]['min'])))
        table_temperature.cell(2, 1).paragraphs[0].alignment = 1
        table_temperature.cell(2, 2).paragraphs[0].add_run(str(round(
            dico_avg_tempe[i]['average'])))
        table_temperature.cell(2, 2).paragraphs[0].alignment = 1
        table_temperature.cell(2, 3).paragraphs[0].add_run(str(round(
            dico_avg_tempe[i]['max'])))
        table_temperature.cell(2, 3).paragraphs[0].alignment = 1

        doc.add_paragraph()

        table_exposure_condition = doc.add_table(
            rows=7, cols=5, style="Table Grid")
        liste_entete = ["Intervention", "Date - Heure",
                        "Température (°C)", "Conductivité (µS/cm)", "pH", "Oxygène dissous (mg/L)"]
        liste_jours = ["J+0", "J+14", "J+N", "J+21"]
        liste_entete_BDD = ["date", "temperature",
                            "conductivity", "ph", "oxygen"]
        for num_entete in range(6):
            table_exposure_condition.cell(num_entete, 0).paragraphs[0].add_run(
                liste_entete[num_entete]).italic = True
            table_exposure_condition.cell(
                num_entete, 0).paragraphs[0].alignment = 1
        for num_jour in range(4):
            table_exposure_condition.cell(0, num_jour+1).paragraphs[0].add_run(
                liste_jours[num_jour]).bold = True
            table_exposure_condition.cell(
                0, num_jour+1).paragraphs[0].alignment = 1
        for num_entete in range(5):
            for num_jour in range(4):
                table_exposure_condition.cell(num_entete+1, num_jour+1).paragraphs[0].add_run(str(
                    dico_exposure_condition[i][liste_jours[num_jour]][liste_entete_BDD[num_entete]]))
                table_exposure_condition.cell(
                    num_entete+1, num_jour+1).paragraphs[0].alignment = 1
        table_exposure_condition.cell(6, 0).merge(
            table_exposure_condition.cell(6, 4))
        table_exposure_condition.cell(
            6, 0).text = "Commentaire : Où le trouver ? "

    doc.save(campaign + "_Rapport_d_expérimentation.docx")
