from . import recuperation_donnee
from docx import Document
from docx.shared import Pt
from docxcompose.composer import Composer
import os
from PIL import Image, ExifTags
import requests
from io import BytesIO


# campaign correspond au nom de la campagne (ex: AG-003-01) et agence est un booléen qui dit si c'est une agence l'eau ou non
def word_main(campaign, agence, path_photo="Photos", path_output="output"):
    path_ressources = "Ressources/"
    campaign = campaign.upper()
    doc = Document(path_ressources + 'Page_de_garde.docx')
    style = doc.styles['Normal']
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    dico_exposure_condition, dico_avg_tempe, dico_geo_mp, dico_geo_agency, dico_type_biotest = recuperation_donnee(
        campaign)
    print('Données récupérées !')
    liste_reference = list(dico_avg_tempe.keys())
    for reference in liste_reference:
        try:
            doc.add_page_break()
            table_geo_1 = doc.add_table(rows=3, cols=6)

            table_geo_1.cell(1, 1).merge(table_geo_1.cell(1, 2))
            table_geo_1.cell(1, 4).merge(table_geo_1.cell(1, 5))
            table_geo_1.cell(2, 0).merge(table_geo_1.cell(2, 1))
            table_geo_1.cell(2, 2).merge(table_geo_1.cell(2, 5))

            if agence:
                table_geo_2 = doc.add_table(rows=5, cols=4)
                for j in range(0, 5):
                    table_geo_2.cell(j, 0).merge(table_geo_2.cell(j, 1))
                for j in range(0, 1):
                    table_geo_2.cell(j, 2).merge(table_geo_2.cell(j, 3))
            else:
                table_geo_2 = doc.add_table(rows=1, cols=4)
                table_geo_2.cell(0, 0).merge(table_geo_2.cell(0, 1))

            header = table_geo_1.rows[0].cells
            header[0].merge(header[-1])
            if agence:
                case_header = table_geo_1.cell(0, 0).paragraphs[0].add_run(dico_geo_agency[reference]['code'] +
                                                                           " : " + dico_geo_agency[reference]['name'] + "   " + reference)
            else:
                case_header = table_geo_1.cell(0, 0).paragraphs[0].add_run("Point " + reference[-5:-3] + " : " +
                                                                           dico_geo_mp[reference]['name_mp'])
            case_header.bold = True
            table_geo_1.cell(0, 0).paragraphs[0].alignment = 1
            width_table = table_geo_1.cell(0, 0).width

            table_geo_1.cell(1, 0).paragraphs[0].add_run(
                'Commune :').bold = True
            if agence:
                table_geo_1.cell(1, 1).paragraphs[0].add_run(
                    dico_geo_agency[reference]['city'] + "    " + dico_geo_agency[reference]['zipcode'])
                table_geo_1.cell(1, 4).paragraphs[0].add_run(
                    dico_geo_agency[reference]['stream'])
            else:
                table_geo_1.cell(1, 1).paragraphs[0].add_run(
                    dico_geo_mp[reference]['city'] + "    " + dico_geo_mp[reference]['zipcode'])
                table_geo_1.cell(1, 4).paragraphs[0].add_run(
                    dico_geo_mp[reference]['stream'])
            table_geo_1.cell(1, 3).paragraphs[0].add_run(
                "Cours d'eau : ").bold = True

            table_geo_1.cell(2, 0).paragraphs[0].add_run(
                "Biotests :").bold = True
            biotest_francais = traduction_type_biotest(
                dico_type_biotest[reference]['biotest'])
            table_geo_1.cell(2, 2).paragraphs[0].add_run(
                biotest_francais)

            if agence:
                table_geo_2.cell(0, 0).paragraphs[0].add_run(
                    "Réseau de surveillance :").bold = True
                table_geo_2.cell(0, 2).paragraphs[0].add_run(
                    dico_geo_agency[reference]['network'])

                table_geo_2.cell(1, 0).paragraphs[0].add_run(
                    "Type d'hydroécorégion :").bold = True
                table_geo_2.cell(1, 2).paragraphs[0].add_run(
                    dico_geo_agency[reference]['hydroecoregion'])

                table_geo_2.cell(2, 0).paragraphs[0].add_run(
                    "Coordonnées Agence Lambert 93 :").bold = True
                table_geo_2.cell(2, 2).paragraphs[0].add_run('Y ' +
                                                             dico_geo_agency[reference]['lambertY'].replace('.', ','))
                table_geo_2.cell(2, 3).paragraphs[0].add_run('X ' +
                                                             dico_geo_agency[reference]['lambertX'].replace('.', ','))

                table_geo_2.cell(3, 0).paragraphs[0].add_run(
                    "Coordonnées BIOMÆ en degrés décimaux : ").bold = True
                longitude = dico_geo_mp[reference]['longitudeSpotted']
                latitude = dico_geo_mp[reference]['latitudeSpotted']
                if longitude == None:
                    longitude = dico_geo_agency[reference]['longitudeTh']
                if latitude == None:
                    latitude = dico_geo_agency[reference]['latitudeTh']
                table_geo_2.cell(3, 2).paragraphs[0].add_run(
                    str(longitude))
                table_geo_2.cell(3, 3).paragraphs[0].add_run(
                    str(latitude))

                table_geo_2.cell(4, 0).paragraphs[0].add_run(
                    "Coordonnées BIOMÆ Lambert 93 : ").bold = True
                table_geo_2.cell(4, 2).paragraphs[0].add_run('Y ' +
                                                             dico_geo_mp[reference]['lambertYSpotted'].replace('.', ','))
                table_geo_2.cell(4, 3).paragraphs[0].add_run('X ' +
                                                             dico_geo_mp[reference]['lambertXSpotted'].replace('.', ','))

            else:
                table_geo_2.cell(0, 0).paragraphs[0].add_run(
                    "Coordonnées BIOMÆ en degrés décimaux : ").bold = True
                longitude = dico_geo_mp[reference]['longitudeSpotted']
                latitude = dico_geo_mp[reference]['latitudeSpotted']
                if longitude == None:
                    longitude = dico_geo_mp[reference]['longitudeTh']
                if latitude == None:
                    latitude = dico_geo_mp[reference]['latitudeTh']
                table_geo_2.cell(0, 2).paragraphs[0].add_run(
                    str(dico_geo_mp[reference]['longitudeSpotted']))
                table_geo_2.cell(0, 3).paragraphs[0].add_run(
                    str(dico_geo_mp[reference]['latitudeSpotted']))

            table_carte = doc.add_table(rows=4, cols=1)
            lon = str(dico_geo_mp[reference]['longitudeSpotted'])
            lat = str(dico_geo_mp[reference]['latitudeSpotted'])
            if (lon != "None") & (lat != "None"):
                access_token = "pk.eyJ1IjoiamJyb25uZXIiLCJhIjoiY2s2cW5kOWQwMHBybjNtcW8yMXJuYmo3aiJ9.z8Ekf7a0RGTZ4jrbJVpq8g"
                # layer = '{"id":"water","source":{"url":"mapbox://mapbox.mapbox-streets-v8","type":"vector"},"source-layer":"water","type":"fill","paint":{"fill-color":"%2300ffff"}}'
                url_street = f"https://api.mapbox.com/styles/v1/mapbox/streets-v11/static/pin-s+FF0000({lon},{lat})/{lon},{lat},9.21/450x300@2x?access_token={access_token}"
                response = requests.get(url_street)
                carte_street = BytesIO(response.content)
                table_carte.cell(0, 0).paragraphs[0].add_run().add_picture(
                    carte_street, width=4500000)
            else:
                table_carte.cell(0, 0).paragraphs[0].add_run().add_picture(
                    path_ressources + "/carre_blanc.jpg", width=4500000)
            table_carte.cell(0, 0).paragraphs[0].alignment = 1
            table_carte.cell(1, 0).text = "Localisation du point de mesure"
            table_carte.cell(1, 0).paragraphs[0].alignment = 1

            if (lon != "None") & (lat != "None"):
                url_satellite = f"https://api.mapbox.com/styles/v1/mapbox/satellite-streets-v11/static/pin-s+FF0000({lon},{lat})/{lon},{lat},13.5/450x300@2x?access_token={access_token}"
                response = requests.get(url_satellite)
                carte_satellite = BytesIO(response.content)
                table_carte.cell(2, 0).paragraphs[0].add_run().add_picture(
                    carte_satellite, width=4500000)
            else:
                table_carte.cell(2, 0).paragraphs[0].add_run().add_picture(
                    path_ressources + "/carre_blanc.jpg", width=4500000)
            table_carte.cell(2, 0).paragraphs[0].alignment = 1
            table_carte.cell(
                3, 0).text = "Vue satellitaire"
            table_carte.cell(3, 0).paragraphs[0].alignment = 1

            doc.add_page_break()

            table_image = doc.add_table(rows=8, cols=2)
            table_image.cell(0, 0).merge(table_image.cell(0, 1))
            table_image.cell(1, 0).merge(table_image.cell(1, 1))

            if agence:
                table_image.cell(0, 0).paragraphs[0].add_run(dico_geo_agency[reference]['code'] +
                                                             " : " + dico_geo_agency[reference]['name']).bold = True
            else:
                table_image.cell(0, 0).paragraphs[0].add_run("Point " + reference[-5:-3] + " : " +
                                                             dico_geo_mp[reference]['name_mp']).bold = True
            table_image.cell(0, 0).paragraphs[0].alignment = 1
            table_image.cell(
                0, 0).paragraphs[0].paragraph_format.line_spacing = font.size

            table_image.cell(1, 0).paragraphs[0].add_run(
                "Photos de la station de mesure de la qualité des eaux pour la campagne " + campaign[-2:] + "-" + dico_exposure_condition[reference]["J+0"]["date"][6:10]).bold = True  # Mettre que l'année, passage en argument ou autre méthode de récupération ?
            table_image.cell(1, 0).paragraphs[0].alignment = 1
            table_image.cell(
                1, 0).paragraphs[0].paragraph_format.line_spacing = font.size

            table_image.cell(3, 0).text = "Aval de zone d’encagement"
            table_image.cell(3, 0).paragraphs[0].alignment = 1
            table_image.cell(3, 1).text = "Amont de zone d’encagement"
            table_image.cell(3, 1).paragraphs[0].alignment = 1
            table_image.cell(5, 0).text = "Gros plan de l’encagement"
            table_image.cell(5, 0).paragraphs[0].alignment = 1
            table_image.cell(5, 1).text = "Panorama encagement"
            table_image.cell(5, 1).paragraphs[0].alignment = 1

            for row in range(3, 8):
                if row != 4:
                    table_image.cell(
                        row, 0).paragraphs[0].paragraph_format.line_spacing = font.size
                    table_image.cell(
                        row, 1).paragraphs[0].paragraph_format.line_spacing = font.size

            nom_photo = recuperation_photo(
                reference, path_photo, path_ressources)
            rotation_image(nom_photo['amont'])
            rotation_image(nom_photo['aval'])
            rotation_image(nom_photo['zoom'])
            rotation_image(nom_photo['panorama'])

            table_image.cell(2, 0).paragraphs[0].add_run().add_picture(
                nom_photo['aval'], width=3046870)  # width=3046870, height=2111370
            table_image.cell(2, 1).paragraphs[0].add_run().add_picture(
                nom_photo['amont'], width=3046870)
            table_image.cell(4, 0).paragraphs[0].add_run().add_picture(
                nom_photo['zoom'], width=3046870)
            table_image.cell(4, 1).paragraphs[0].add_run().add_picture(
                nom_photo['panorama'], width=3046870)
            for elt in [(2, 0), (2, 1), (4, 0), (4, 1)]:
                table_image.cell(elt[0],
                                 elt[1]).paragraphs[0].paragraph_format.space_after = Pt(0)
                table_image.cell(elt[0],
                                 elt[1]).paragraphs[0].paragraph_format.space_before = Pt(0)

            table_image.cell(6, 0).paragraphs[0].add_run(
                "Type de système d’exposition : ").bold = True

            type_barrel_J0 = dico_exposure_condition[reference]['J+0']['type']
            if (type_barrel_J0 == 'barrel'):
                type_barrel_J0 = 'Fut'
            elif (type_barrel_J0 == 'box'):
                type_barrel_J0 = 'Caisse'
            table_image.cell(6, 1).paragraphs[0].add_run(type_barrel_J0)
            table_image.cell(7, 0).merge(table_image.cell(7, 1))
            table_image.cell(7, 0).paragraphs[0].add_run(
                "Paramètres physico-chimiques pour la campagne : " + campaign[-2:] + "-" + dico_exposure_condition[reference]["J+0"]["date"][6:10]).bold = True

            table_temperature = doc.add_table(
                rows=2, cols=4, style="Table Grid")
            table_temperature.cell(0, 0).merge(table_temperature.cell(1, 0))
            table_temperature.cell(0, 0).paragraphs[0].add_run(
                "Température eau (°C) Sonde en continu").italic = True
            table_temperature.cell(0, 0).paragraphs[0].alignment = 1

            table_temperature.cell(0, 1).paragraphs[0].add_run(
                "Minimum")  # .bold = True
            table_temperature.cell(0, 1).paragraphs[0].alignment = 1
            table_temperature.cell(0, 2).paragraphs[0].add_run(
                "Moyenne")  # .bold = True
            table_temperature.cell(0, 2).paragraphs[0].alignment = 1
            table_temperature.cell(0, 3).paragraphs[0].add_run(
                "Maximum")  # .bold = True
            table_temperature.cell(0, 3).paragraphs[0].alignment = 1
            if dico_avg_tempe[reference]['min'] is not None:
                table_temperature.cell(1, 1).paragraphs[0].add_run(str(round(
                    dico_avg_tempe[reference]['min'], 1)).replace('.', ','))
            table_temperature.cell(1, 1).paragraphs[0].alignment = 1
            if dico_avg_tempe[reference]['average'] is not None:
                table_temperature.cell(1, 2).paragraphs[0].add_run(str(round(
                    dico_avg_tempe[reference]['average'], 1)).replace('.', ','))
            table_temperature.cell(1, 2).paragraphs[0].alignment = 1
            if dico_avg_tempe[reference]['max'] is not None:
                table_temperature.cell(1, 3).paragraphs[0].add_run(str(round(
                    dico_avg_tempe[reference]['max'], 1)).replace('.', ','))
            table_temperature.cell(1, 3).paragraphs[0].alignment = 1
            for row in range(2):
                for col in range(4):
                    paragraph = table_temperature.cell(
                        row, col).paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.space_before = Pt(4)

            interligne = doc.add_paragraph()
            interligne.paragraph_format.space_after = Pt(0)
            interligne.paragraph_format.space_before = Pt(0)

            chemistry = 1 if "chemistry" in dico_type_biotest[reference]['biotest'] else 0
            if dico_exposure_condition[reference]['fusion?']:
                liste_jours = ["J+0", "J+14", "J+N", "J+21"]
            else:
                liste_jours = ["J+0", "J+7", "J+N", "J+21"]
            liste_indice_jours_utiles = []
            for num_jour in range(4):
                if dico_exposure_condition[reference][liste_jours[num_jour]]['date'] != None:
                    liste_indice_jours_utiles.append(num_jour)
            nombre_jours_utiles = len(liste_indice_jours_utiles)
            table_exposure_condition = doc.add_table(
                rows=7+chemistry, cols=1+nombre_jours_utiles, style="Table Grid")
            liste_entete = ["Intervention", "Date - Heure",
                            "Température (°C)", "Conductivité (µS/cm)", "pH", "Oxygène dissous (mg/L)", "Survie Chimie (%)"]
            liste_entete_BDD = ["date", "temperature",
                                "conductivity", "ph", "oxygen"]
            for num_entete in range(6+chemistry):
                paragraph = table_exposure_condition.cell(
                    num_entete, 0).paragraphs[0]
                paragraph.add_run(liste_entete[num_entete]).italic = True
                paragraph.alignment = 1
                table_exposure_condition.cell(
                    num_entete, 0).width = width_table*0.3
            for num_jour in range(nombre_jours_utiles):
                paragraph = table_exposure_condition.cell(
                    0, num_jour+1).paragraphs[0]
                paragraph.add_run(
                    liste_jours[liste_indice_jours_utiles[num_jour]]).bold = True
                paragraph.alignment = 1
            for num_entete in range(5):
                for num_jour in range(nombre_jours_utiles):
                    paragraph = table_exposure_condition.cell(
                        num_entete+1, num_jour+1).paragraphs[0]
                    value = dico_exposure_condition[reference][liste_jours[liste_indice_jours_utiles[num_jour]]
                                                               ][liste_entete_BDD[num_entete]]
                    if (liste_entete_BDD[num_entete] == "conductivity") & (value is not None):
                        paragraph.add_run(str(round(
                            value)).replace('.', ','))
                    else:
                        paragraph.add_run(str(
                            value).replace('.', ','))

                    paragraph.alignment = 1

            if chemistry:
                table_exposure_condition.cell(6, 1).merge(
                    table_exposure_condition.cell(6, nombre_jours_utiles))
                paragraph_survie = table_exposure_condition.cell(
                    6, 1).paragraphs[0]
                paragraph_survie.add_run(str(
                    round(dico_type_biotest[reference]['survivor_chemistry'], 1)).replace('.', ','))
                paragraph_survie.alignment = 1

                table_exposure_condition.cell(7, 0).merge(
                    table_exposure_condition.cell(7, nombre_jours_utiles))
                paragraph_comment = table_exposure_condition.cell(
                    7, 0).paragraphs[0]
            else:
                table_exposure_condition.cell(6, 0).merge(
                    table_exposure_condition.cell(6, nombre_jours_utiles))
                paragraph_comment = table_exposure_condition.cell(
                    6, 0).paragraphs[0]

            comment = ""
            for jour in liste_jours:
                if dico_exposure_condition[reference][jour]['comment'] != None:
                    comment += f"{jour} : {dico_exposure_condition[reference][jour]['comment']}\n"
            paragraph_comment.add_run(comment[:-1])
            for num_entete in range(7+chemistry):
                for num_jour in range(1+nombre_jours_utiles):
                    paragraph = table_exposure_condition.cell(
                        num_entete, num_jour).paragraphs[0]
                    paragraph.paragraph_format.space_after = Pt(4)
                    paragraph.paragraph_format.space_before = Pt(4)
            print(f'Page de la référence {reference} créée ! :D')
        except Exception as e:
            print(reference + " n'a pas été créée")
            print(e)
            pass

    doc.add_page_break()

    composer = Composer(doc)
    page_fin = Document(path_ressources + 'Page_fin.docx')
    composer.append(page_fin)
    name_doc = campaign + "_Rapport_d_expérimentation.docx"
    composer.save(path_output + "/" + name_doc)


def traduction_type_biotest(biotest_anglais):
    biotest_francais = []
    for elt in biotest_anglais:
        if elt == "neurology":
            biotest_francais.append("Neurotoxicité")
        if elt == "alimentation":
            biotest_francais.append("Alimentation")
        if elt == "chemistry":
            biotest_francais.append("Chimie")
        if elt == "reproduction":
            biotest_francais.append("Reproduction")
    string = ""
    for elt in biotest_francais:
        string += elt + ", "
    string = string[:-2]
    return string


def recuperation_photo(reference, path_photo, path_ressources):
    prefixe = path_photo + "/" + reference
    list_type = ["amont", "aval", "zoom", "panorama"]
    dico_nom = {}
    try:
        filenames = os.listdir(prefixe)
    except FileNotFoundError:
        for typ in list_type:
            dico_nom[typ] = path_ressources + "/carre_blanc.jpg"
    else:
        for elt in filenames:
            l_nom = elt.split("_")
            type_photo = l_nom[3].lower()
            dico_nom[type_photo] = prefixe + "/" + elt

        for typ in list_type:
            try:
                dico_nom[typ]
            except KeyError:
                dico_nom[typ] = path_ressources + "/carre_blanc.jpg"
    return dico_nom


def rotation_image(path_photo):
    try:
        image = Image.open(path_photo)
        for orientation in ExifTags.TAGS.keys():
            if ExifTags.TAGS[orientation] == 'Orientation':
                break
        exif = dict(image._getexif().items())
        if exif[orientation] == 3:
            image = image.rotate(180, expand=True)
        elif exif[orientation] == 6:
            image = image.rotate(270, expand=True)
        elif exif[orientation] == 8:
            image = image.rotate(90, expand=True)
        image.save(path_photo)
        image.close()

    except (AttributeError, KeyError, IndexError):
        # cases: image don't have getexif
        pass
